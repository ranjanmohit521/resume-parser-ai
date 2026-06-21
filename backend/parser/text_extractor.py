"""
text_extractor.py
-----------------
Converts various file formats (PDF, DOCX, TXT, Images) into plain text.
"""

import os
import io
from pathlib import Path

# ---------------------------------------------------------------------------
# PDF Extraction
# ---------------------------------------------------------------------------
def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from PDF bytes using pdfminer.six with PyPDF2 as fallback."""
    text = ""

    # Primary: pdfminer.six
    try:
        from pdfminer.high_level import extract_text as pdfminer_extract
        from pdfminer.layout import LAParams
        text = pdfminer_extract(io.BytesIO(file_bytes), laparams=LAParams())
        if text and text.strip():
            return text.strip()
    except Exception:
        pass

    # Fallback: PyPDF2
    try:
        import PyPDF2
        reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages)
    except Exception:
        pass

    return text.strip()


# ---------------------------------------------------------------------------
# DOCX Extraction
# ---------------------------------------------------------------------------
def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from DOCX bytes using python-docx."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        paragraphs = [para.text for para in doc.paragraphs]
        # Also extract table cells
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(paragraphs).strip()
    except Exception as e:
        raise ValueError(f"Failed to extract DOCX: {e}")


# ---------------------------------------------------------------------------
# TXT Extraction
# ---------------------------------------------------------------------------
def extract_text_from_txt(file_bytes: bytes) -> str:
    """Decode plain text from bytes."""
    for encoding in ("utf-8", "latin-1", "cp1252"):
        try:
            return file_bytes.decode(encoding).strip()
        except UnicodeDecodeError:
            continue
    return file_bytes.decode("utf-8", errors="replace").strip()


# ---------------------------------------------------------------------------
# Image Extraction (OCR via Tesseract)
# ---------------------------------------------------------------------------
def extract_text_from_image(file_bytes: bytes) -> str:
    """Use pytesseract OCR to extract text from image bytes."""
    try:
        import pytesseract
        from PIL import Image

        image = Image.open(io.BytesIO(file_bytes))
        # Convert to RGB if needed (handles RGBA/grayscale)
        if image.mode not in ("RGB", "L"):
            image = image.convert("RGB")

        text = pytesseract.image_to_string(image, lang="eng")
        return text.strip()
    except ImportError:
        raise ValueError(
            "pytesseract or Pillow is not installed. "
            "Install with: pip install pytesseract Pillow\n"
            "Also install Tesseract OCR: https://github.com/UB-Mannheim/tesseract/wiki"
        )
    except Exception as e:
        raise ValueError(f"Image OCR failed: {e}")


# ---------------------------------------------------------------------------
# Main Dispatcher
# ---------------------------------------------------------------------------
SUPPORTED_EXTENSIONS = {
    ".pdf": extract_text_from_pdf,
    ".docx": extract_text_from_docx,
    ".txt": extract_text_from_txt,
    ".jpg": extract_text_from_image,
    ".jpeg": extract_text_from_image,
    ".png": extract_text_from_image,
    ".bmp": extract_text_from_image,
    ".tiff": extract_text_from_image,
    ".tif": extract_text_from_image,
}


def extract_text(file_bytes: bytes, filename: str) -> str:
    """
    Dispatch text extraction based on file extension.

    Args:
        file_bytes: Raw bytes of the uploaded file.
        filename:   Original filename (used to detect format).

    Returns:
        Extracted plain text string.

    Raises:
        ValueError: If the file format is unsupported.
    """
    ext = Path(filename).suffix.lower()
    extractor = SUPPORTED_EXTENSIONS.get(ext)

    if extractor is None:
        raise ValueError(
            f"Unsupported file format: '{ext}'. "
            f"Supported formats: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
        )

    return extractor(file_bytes)
